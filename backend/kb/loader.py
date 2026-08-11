"""
文档加载器
=========

支持 PDF（PyMuPDF）和 Word（python-docx）格式的文档解析。
输出统一的 ParsedDocument 格式，包含章节结构和页码信息。

PDF 解析策略:
    - 使用 PyMuPDF (fitz) 逐页提取文本
    - 基于正则识别章节标题（"第X章"、"X.X" 等模式）
    - 保留页码信息用于后续引用
    - 识别 LaTeX 公式区域

Word 解析策略:
    - 使用 python-docx 读取段落和样式
    - 基于 Heading 样式识别章节层级
    - 处理公式对象
"""

import os
import re
import logging
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
from pathlib import Path

logger = logging.getLogger(__name__)


# ==================== 数据结构 ====================

@dataclass
class ContentElement:
    """文档内容最小单元"""
    type: str           # "paragraph" | "heading" | "theorem" | "proof" |
                        # "definition" | "example" | "formula"
    text: str           # 纯文本内容
    page_number: int
    element_id: str     # 唯一标识


@dataclass
class Chapter:
    """章节结构"""
    title: str
    level: int           # 1=章, 2=节, 3=小节
    heading_text: str
    page_start: int
    elements: List[ContentElement] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """解析后的统一文档格式"""
    source_path: str
    source_filename: str
    total_pages: int
    chapters: List[Chapter] = field(default_factory=list)
    # 未被分配到任何章节的内容（如前言、目录）
    unassigned_elements: List[ContentElement] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)


# ==================== 章节标题识别 ====================

# 中文章节标题模式
CHAPTER_PATTERNS = [
    (re.compile(r'^第([一二三四五六七八九十\d]+)章\s+(.+)'), 1),       # 第X章 XXX
    (re.compile(r'^第([一二三四五六七八九十\d]+)节\s+(.+)'), 2),       # 第X节 XXX
    (re.compile(r'^(\d+(?:\.\d+)*)\s+(.+)'), None),                    # 1.1 XXX, 1.1.1 XXX
]

# 定理环境关键词
THEOREM_KEYWORDS = {
    'theorem':   re.compile(r'^(定理\s*\d|定理\s*[（(])'),
    'definition': re.compile(r'^(定义\s*\d|定义\s*[（(])'),
    'proof':     re.compile(r'^(证明[：:]|证明\s*$|Proof)'),
    'example':   re.compile(r'^(例\s*\d|例\s*[（(]|例题\s*\d)'),
    'axiom':     re.compile(r'^(公理\s*\d|公理\s*[（(])'),
    'corollary': re.compile(r'^(推论\s*\d|推论\s*[（(])'),
    'lemma':     re.compile(r'^(引理\s*\d|引理\s*[（(])'),
}


def detect_element_type(text: str) -> str:
    """检测文本块的内容类型"""
    text_stripped = text.strip()
    for elem_type, pattern in THEOREM_KEYWORDS.items():
        if pattern.search(text_stripped):
            return elem_type
    return "paragraph"


def detect_chapter_level(heading_text: str) -> Tuple[Optional[int], str]:
    """
    检测标题层级。返回 (level, title)。
    level: 1=章, 2=节, 3=小节, None=非章节标题
    """
    for pattern, level in CHAPTER_PATTERNS:
        m = pattern.match(heading_text.strip())
        if m:
            if level is not None:
                return level, heading_text.strip()
            else:
                # 数字编号: 1.1 → level 2, 1.1.1 → level 3
                num_parts = m.group(1).split('.')
                return min(len(num_parts) + 1, 3), heading_text.strip()
    return None, heading_text.strip()


# ==================== 文档加载器 ====================

class DocumentLoader:
    """文档加载器：支持 PDF 和 Word 格式"""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md'}

    def __init__(self):
        self._loaders = {
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.doc': self._load_docx,
            '.txt': self._load_text,
            '.md': self._load_text,
        }

    def load(self, file_path: str) -> ParsedDocument:
        """加载文档，自动识别格式"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}。支持: {self.SUPPORTED_EXTENSIONS}")

        loader = self._loaders.get(ext)
        if loader is None:
            raise ValueError(f"未找到 {ext} 格式的加载器")

        logger.info(f"开始加载文档: {path.name}")
        doc = loader(str(path))
        doc.source_path = str(path)
        doc.source_filename = path.name
        logger.info(f"文档加载完成: {path.name}, {doc.total_pages}页, "
                     f"{len(doc.chapters)}个章节, {len(doc.errors)}个错误")
        return doc

    # ---------- PDF 解析 ----------

    def _load_pdf(self, file_path: str) -> ParsedDocument:
        """使用 PyMuPDF 解析 PDF"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("请安装 PyMuPDF: pip install pymupdf")

        doc = ParsedDocument(
            source_path=file_path,
            source_filename=Path(file_path).name,
            total_pages=0,
        )

        try:
            pdf = fitz.open(file_path)
            doc.total_pages = pdf.page_count
        except Exception as e:
            doc.errors.append(f"无法打开 PDF: {e}")
            return doc

        current_chapter: Optional[Chapter] = None
        element_counter = 0
        full_text_pages = []  # 按页存储全文，用于后续章节分析

        # 第一遍：逐页提取文本
        for page_num in range(pdf.page_count):
            try:
                page = pdf[page_num]
                text = page.get_text("text")
                if text.strip():
                    full_text_pages.append((page_num + 1, text))
            except Exception as e:
                doc.errors.append(f"第{page_num+1}页提取失败: {e}")

        # 第二遍：解析段落结构
        current_text_buffer = []
        current_page = 1

        for page_num, page_text in full_text_pages:
            lines = page_text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    # 空行可能是段落边界，刷新缓冲区
                    if current_text_buffer:
                        self._emit_paragraph(
                            doc, current_text_buffer, current_page, element_counter
                        )
                        element_counter += 1
                        current_text_buffer = []
                    continue

                # 检查是否是章节标题
                level, title = detect_chapter_level(line)
                if level is not None:
                    # 先保存之前的缓冲区
                    if current_text_buffer:
                        self._emit_paragraph(
                            doc, current_text_buffer, current_page, element_counter
                        )
                        element_counter += 1
                        current_text_buffer = []

                    # 创建新章节
                    current_chapter = Chapter(
                        title=title,
                        level=level,
                        heading_text=line,
                        page_start=page_num,
                    )
                    doc.chapters.append(current_chapter)
                    continue

                current_text_buffer.append(line)
                current_page = page_num

        # 处理最后的缓冲区
        if current_text_buffer:
            self._emit_paragraph(doc, current_text_buffer, current_page, element_counter)

        pdf.close()
        return doc

    def _emit_paragraph(self, doc: ParsedDocument, lines: List[str],
                        page: int, counter: int) -> ContentElement:
        """将缓冲的行组装为一个 ContentElement"""
        text = ' '.join(lines)
        elem_type = detect_element_type(text)
        elem = ContentElement(
            type=elem_type,
            text=text,
            page_number=page,
            element_id=f"elem_{counter:05d}",
        )
        # 添加到当前章节或未分配列表
        if doc.chapters:
            doc.chapters[-1].elements.append(elem)
        else:
            doc.unassigned_elements.append(elem)
        return elem

    # ---------- Word (docx) 解析 ----------

    def _load_docx(self, file_path: str) -> ParsedDocument:
        """使用 python-docx 解析 Word 文档"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = ParsedDocument(
            source_path=file_path,
            source_filename=Path(file_path).name,
            total_pages=0,  # python-docx 无法直接获取页数
        )

        try:
            docx = Document(file_path)
        except Exception as e:
            doc.errors.append(f"无法打开 Word 文档: {e}")
            return doc

        current_chapter: Optional[Chapter] = None
        element_counter = 0
        page_estimate = 1  # 估算页码：每~50段为1页

        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""

            # 检测标题样式（Word 中的 Heading 1/2/3）
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 1
                current_chapter = Chapter(
                    title=text,
                    level=min(level, 3),
                    heading_text=text,
                    page_start=page_estimate,
                )
                doc.chapters.append(current_chapter)
                continue

            # 检测非样式标题（纯文本但符合章节模式）
            level, title = detect_chapter_level(text)
            if level is not None:
                current_chapter = Chapter(
                    title=title,
                    level=level,
                    heading_text=text,
                    page_start=page_estimate,
                )
                doc.chapters.append(current_chapter)
                continue

            # 普通段落
            elem_type = detect_element_type(text)
            elem = ContentElement(
                type=elem_type,
                text=text,
                page_number=page_estimate,
                element_id=f"elem_{element_counter:05d}",
            )
            element_counter += 1

            if current_chapter:
                current_chapter.elements.append(elem)
            else:
                doc.unassigned_elements.append(elem)

            # 每 50 段估算翻页
            if element_counter % 50 == 0:
                page_estimate += 1

        doc.total_pages = page_estimate
        return doc

    # ---------- 纯文本解析 ----------

    def _load_text(self, file_path: str) -> ParsedDocument:
        """解析纯文本或 Markdown 文件"""
        doc = ParsedDocument(
            source_path=file_path,
            source_filename=Path(file_path).name,
            total_pages=1,
        )

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='gbk') as f:
                content = f.read()

        current_chapter: Optional[Chapter] = None
        element_counter = 0
        lines = content.split('\n')

        # Markdown 标题
        md_heading = re.compile(r'^(#{1,6})\s+(.+)')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Markdown 标题
            m = md_heading.match(line)
            if m:
                level = min(len(m.group(1)), 3)
                title = m.group(2)
                current_chapter = Chapter(
                    title=title,
                    level=level,
                    heading_text=line,
                    page_start=1,
                )
                doc.chapters.append(current_chapter)
                continue

            # 中文章节标题
            level, title = detect_chapter_level(line)
            if level is not None:
                current_chapter = Chapter(
                    title=title,
                    level=level,
                    heading_text=line,
                    page_start=1,
                )
                doc.chapters.append(current_chapter)
                continue

            elem = ContentElement(
                type=detect_element_type(line),
                text=line,
                page_number=1,
                element_id=f"elem_{element_counter:05d}",
            )
            element_counter += 1

            if current_chapter:
                current_chapter.elements.append(elem)
            else:
                doc.unassigned_elements.append(elem)

        return doc
